import io
from typing import Union
from docx import Document
from PyPDF2 import PdfReader

def load_resume_text(uploaded_file: Union[io.BytesIO, object, str]) -> str:
    """
    Accepts a Streamlit uploaded_file (file-like) or path string.
    Returns extracted plain text.
    """
    # If path string
    if isinstance(uploaded_file, str):
        path = uploaded_file
        name = os.path.basename(path)
        open_mode = True
    else:
        # Streamlit uploaded file - has .name and .read()
        name = getattr(uploaded_file, "name", "uploaded")
        open_mode = False

    text = ""

    if name.lower().endswith(".pdf"):
        # PdfReader accepts file-like objects
        if open_mode:
            reader = PdfReader(path)
        else:
            reader = PdfReader(uploaded_file)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"

    elif name.lower().endswith(".docx"):
        # python-docx Document can accept a path-like or file-like
        if open_mode:
            doc = Document(path)
        else:
            # write to temp bytes as Document() can't accept BytesIO with some versions
            tmp = uploaded_file.read()
            doc = Document(io.BytesIO(tmp))
        for p in doc.paragraphs:
            text += p.text + "\n"

    elif name.lower().endswith(".txt"):
        if open_mode:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            # uploaded file: read bytes and decode
            text = uploaded_file.read().decode("utf-8", errors="ignore")

    else:
        raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")

    # basic cleanup
    text = " ".join(text.split())
    return text
